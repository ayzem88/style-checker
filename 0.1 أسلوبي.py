# -*- coding: utf-8 -*-
"""
أسلوبي - النسخة الرسومية
أداة متطورة لتصحيح الأخطاء اللغوية في النصوص العربية
واجهة رسومية بألوان رمادية
"""

import sys
import os
import re
import json
import subprocess
from pathlib import Path
from datetime import datetime

# محاولة استيراد المكتبات المطلوبة
try:
    from PyQt6.QtWidgets import (
        QApplication,
        QMainWindow,
        QWidget,
        QVBoxLayout,
        QHBoxLayout,
        QPushButton,
        QTextEdit,
        QLabel,
        QFileDialog,
        QMessageBox,
        QSplitter,
        QToolBar,
        QStatusBar,
        QMenuBar,
        QMenu,
        QSizePolicy,
        QListWidget,
        QListWidgetItem,
    )
    from PyQt6.QtCore import Qt, QSize
    from PyQt6.QtGui import (
        QFont,
        QColor,
        QPalette,
        QAction,
        QKeySequence,
        QTextCharFormat,
        QTextCursor,
    )

    PYQT6_AVAILABLE = True
except ImportError:
    PYQT6_AVAILABLE = False
    print("❌ PyQt6 غير مثبت. قم بتثبيته باستخدام: pip install PyQt6")

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import RGBColor, Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ألوان رمادية
COLORS = {
    "bg_main": "#e8e8e8",  # الخلفية الرئيسية (رمادي فاتح)
    "bg_secondary": "#d0d0d0",  # الخلفية الثانوية
    "bg_tertiary": "#f5f5f5",  # الخلفية الثالثية (رمادي فاتح جداً)
    "fg_primary": "#2d2d2d",  # النص الأساسي (داكن)
    "fg_secondary": "#5a5a5a",  # النص الثانوي
    "accent": "#6b6b6b",  # اللون المميز (رمادي متوسط)
    "border": "#a0a0a0",  # الحدود
    "hover": "#b8b8b8",  # عند التمرير
    "toolbar": "#d8d8d8",  # شريط الأدوات
    "error": "#8b0000",  # لون الأخطاء
    "success": "#006400",  # لون النجاح
    "error_bg": "#ffcccc",  # خلفية الكلمة الخاطئة (أحمر فاتح)
    "success_bg": "#ccffcc",  # خلفية الكلمة الصحيحة (أخضر فاتح)
}


class StyleCheckerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.corrections = {}
        self.wrong_words = set()
        self.errors = []
        self.corrections_file_path = None

        # البحث عن ملف التصحيحات تلقائياً
        self.find_corrections_file()

        self.init_ui()
        self.setup_gray_theme()

    def find_corrections_file(self):
        """البحث عن ملف التصحيحات في المجلد الحالي"""
        current_dir = os.path.dirname(os.path.abspath(__file__))
        corrections_path = os.path.join(current_dir, "corrections.json")

        if os.path.exists(corrections_path):
            self.corrections_file_path = corrections_path
            self.load_corrections(corrections_path)
        else:
            # البحث في المجلدات الفرعية
            for root, dirs, files in os.walk(current_dir):
                if "corrections.json" in files:
                    self.corrections_file_path = os.path.join(root, "corrections.json")
                    self.load_corrections(self.corrections_file_path)
                    break

    def load_corrections(self, file_path):
        """تحميل ملف التصحيحات"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                self.corrections = json.load(file)
            self.wrong_words = set(self.corrections.keys())
            self.statusBar().showMessage(
                f"تم تحميل {len(self.corrections)} تصحيح", 3000
            )
            return True
        except Exception as e:
            QMessageBox.warning(self, "تحذير", f"فشل في تحميل ملف التصحيحات:\n{str(e)}")
            return False

    def init_ui(self):
        """تهيئة الواجهة"""
        self.setWindowTitle("أسلوبي - أداة تصحيح الأخطاء اللغوية")
        self.setGeometry(100, 100, 1400, 800)

        # إنشاء القائمة الرئيسية
        self.create_menu_bar()

        # إنشاء شريط الأدوات
        self.create_toolbar()

        # إنشاء المنطقة الرئيسية
        self.create_main_area()

        # إنشاء شريط الحالة مع الإحصائيات
        self.statusBar().showMessage("جاهز")
        self.stats_label = QLabel("عدد الأخطاء: 0")
        self.stats_label.setStyleSheet(
            f"color: {COLORS['fg_secondary']}; padding: 5px; font-size: 11px;"
        )
        self.statusBar().addPermanentWidget(self.stats_label)

    def setup_gray_theme(self):
        """إعداد الثيم الرمادي"""
        palette = QPalette()

        # الألوان الأساسية
        palette.setColor(QPalette.ColorRole.Window, QColor(COLORS["bg_main"]))
        palette.setColor(QPalette.ColorRole.WindowText, QColor(COLORS["fg_primary"]))
        palette.setColor(QPalette.ColorRole.Base, QColor(COLORS["bg_tertiary"]))
        palette.setColor(
            QPalette.ColorRole.AlternateBase, QColor(COLORS["bg_secondary"])
        )
        palette.setColor(QPalette.ColorRole.ToolTipBase, QColor(COLORS["bg_secondary"]))
        palette.setColor(QPalette.ColorRole.ToolTipText, QColor(COLORS["fg_primary"]))
        palette.setColor(QPalette.ColorRole.Text, QColor(COLORS["fg_primary"]))
        palette.setColor(QPalette.ColorRole.Button, QColor(COLORS["bg_tertiary"]))
        palette.setColor(QPalette.ColorRole.ButtonText, QColor(COLORS["fg_primary"]))
        palette.setColor(QPalette.ColorRole.BrightText, QColor("#000000"))
        palette.setColor(QPalette.ColorRole.Link, QColor(COLORS["accent"]))
        palette.setColor(QPalette.ColorRole.Highlight, QColor(COLORS["accent"]))
        palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))

        self.setPalette(palette)

    def create_menu_bar(self):
        """إنشاء شريط القوائم"""
        menubar = self.menuBar()

        # قائمة الملف
        file_menu = menubar.addMenu("الملفات")

        # تحميل ملف التصحيحات
        load_corrections_action = QAction("تحميل ملف التصحيحات", self)
        load_corrections_action.setShortcut(QKeySequence("Ctrl+L"))
        load_corrections_action.triggered.connect(self.load_corrections_file)
        file_menu.addAction(load_corrections_action)

        file_menu.addSeparator()

        # استيراد TXT
        import_txt_action = QAction(" txt استيراد", self)
        import_txt_action.setShortcut(QKeySequence("Ctrl+O"))
        import_txt_action.triggered.connect(self.import_txt)
        file_menu.addAction(import_txt_action)

        # استيراد DOCX
        if DOCX_AVAILABLE:
            import_docx_action = QAction(" word استيراد", self)
            import_docx_action.triggered.connect(self.import_docx)
            file_menu.addAction(import_docx_action)
        else:
            import_docx_action = QAction("استيراد word (غير متاح)", self)
            import_docx_action.setEnabled(False)
            file_menu.addAction(import_docx_action)

        file_menu.addSeparator()

        # تصدير النتيجة
        export_action = QAction("تصدير النتيجة", self)
        export_action.setShortcut(QKeySequence("Ctrl+S"))
        export_action.triggered.connect(self.export_result)
        file_menu.addAction(export_action)

        # تصدير تقرير DOCX
        if DOCX_AVAILABLE:
            export_docx_action = QAction("تصدير تقرير word", self)
            export_docx_action.triggered.connect(self.export_docx_report)
            file_menu.addAction(export_docx_action)

        file_menu.addSeparator()

        # خروج
        exit_action = QAction("خروج", self)
        exit_action.setShortcut(QKeySequence("Ctrl+Q"))
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # قائمة الأدوات
        tools_menu = menubar.addMenu("الأدوات")

        check_action = QAction("فحص النص", self)
        check_action.setShortcut(QKeySequence("Ctrl+Return"))
        check_action.triggered.connect(self.check_text)
        tools_menu.addAction(check_action)

        clear_action = QAction("مسح الكل", self)
        clear_action.setShortcut(QKeySequence("Ctrl+Del"))
        clear_action.triggered.connect(self.clear_all)
        tools_menu.addAction(clear_action)

    def create_toolbar(self):
        """إنشاء شريط الأدوات"""
        toolbar = QToolBar("شريط الأدوات")
        toolbar.setStyleSheet(f"background-color: {COLORS['toolbar']};")
        self.addToolBar(toolbar)

        # Spacer على اليسار
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        toolbar.addWidget(spacer)

        # قسم الأدوات (يبدأ من اليمين)
        tools_label = QLabel("")
        tools_label.setStyleSheet(f"color: {COLORS['fg_secondary']}; padding: 5px;")
        toolbar.addWidget(tools_label)

        btn_clear = QPushButton("مسح الكل")
        btn_clear.setStyleSheet(self.get_button_style())
        btn_clear.clicked.connect(self.clear_all)
        toolbar.addWidget(btn_clear)

        btn_check = QPushButton("فحص النص")
        btn_check.setStyleSheet(self.get_button_style())
        btn_check.clicked.connect(self.check_text)
        toolbar.addWidget(btn_check)

        toolbar.addSeparator()

        # قسم الملفات
        if DOCX_AVAILABLE:
            btn_export_docx = QPushButton("تصدير تقرير")
            btn_export_docx.setStyleSheet(self.get_button_style())
            btn_export_docx.clicked.connect(self.export_docx_report)
            toolbar.addWidget(btn_export_docx)

        btn_export = QPushButton("تصدير النتيجة")
        btn_export.setStyleSheet(self.get_button_style())
        btn_export.clicked.connect(self.export_result)
        toolbar.addWidget(btn_export)

        if DOCX_AVAILABLE:
            btn_import_docx = QPushButton("استيراد word")
            btn_import_docx.setStyleSheet(self.get_button_style())
            btn_import_docx.clicked.connect(self.import_docx)
            toolbar.addWidget(btn_import_docx)

        btn_import_txt = QPushButton("استيراد txt")
        btn_import_txt.setStyleSheet(self.get_button_style())
        btn_import_txt.clicked.connect(self.import_txt)
        toolbar.addWidget(btn_import_txt)

        file_label = QLabel("")
        file_label.setStyleSheet(f"color: {COLORS['fg_secondary']}; padding: 5px;")
        toolbar.addWidget(file_label)

    def get_button_style(self):
        """إرجاع نمط الأزرار"""
        return f"""
            QPushButton {{
                background-color: {COLORS["bg_tertiary"]};
                color: {COLORS["fg_primary"]};
                border: 1px solid {COLORS["border"]};
                padding: 5px 10px;
                border-radius: 3px;
            }}
            QPushButton:hover {{
                background-color: {COLORS["hover"]};
            }}
            QPushButton:pressed {{
                background-color: {COLORS["bg_secondary"]};
            }}
            QPushButton:disabled {{
                background-color: {COLORS["bg_secondary"]};
                color: {COLORS["fg_secondary"]};
            }}
        """

    def create_main_area(self):
        """إنشاء المنطقة الرئيسية"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(10, 10, 10, 10)

        # Splitter لتقسيم الشاشة
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # المربع الأول - النص المصحح مع inline diffs (على اليسار)
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(5, 5, 5, 5)

        left_label = QLabel("النص المصحح")
        left_label.setStyleSheet(
            f"color: {COLORS['fg_primary']}; font-weight: bold; font-size: 12px; padding: 5px;"
        )
        left_layout.addWidget(left_label)

        # مربع النص المصحح مع التغييرات
        self.corrected_text = QTextEdit()
        self.corrected_text.setReadOnly(True)
        self.corrected_text.setStyleSheet(f"""
            QTextEdit {{
                background-color: {COLORS["bg_tertiary"]};
                color: {COLORS["fg_primary"]};
                border: 1px solid {COLORS["border"]};
                border-radius: 3px;
                padding: 10px;
                font-size: 14px;
            }}
        """)
        self.corrected_text.setFont(QFont("Sakkala Majalla", 14))
        left_layout.addWidget(self.corrected_text)

        # ملخص الأخطاء في الأسفل
        summary_label = QLabel("ملخص الأخطاء")
        summary_label.setStyleSheet(
            f"color: {COLORS['fg_primary']}; font-weight: bold; font-size: 12px; padding: 5px;"
        )
        left_layout.addWidget(summary_label)

        self.errors_list = QListWidget()
        self.errors_list.setLayoutDirection(
            Qt.LayoutDirection.RightToLeft
        )  # اتجاه من اليمين لليسار
        self.errors_list.setStyleSheet(f"""
            QListWidget {{
                background-color: {COLORS["bg_tertiary"]};
                color: {COLORS["fg_primary"]};
                border: 1px solid {COLORS["border"]};
                border-radius: 3px;
                padding: 10px;
                font-size: 12px;
                text-align: right;
            }}
            QListWidget::item {{
                padding: 5px;
                border-bottom: 1px solid {COLORS["border"]};
                text-align: right;
            }}
            QListWidget::item:hover {{
                background-color: {COLORS["hover"]};
            }}
            QListWidget::item:selected {{
                background-color: {COLORS["accent"]};
                color: white;
            }}
        """)
        self.errors_list.setFont(QFont("Sakkala Majalla", 12))
        self.errors_list.setMaximumHeight(150)
        self.errors_list.itemClicked.connect(self.highlight_error_in_text)
        left_layout.addWidget(self.errors_list)

        splitter.addWidget(left_widget)

        # المربع الثاني - النص الأصلي (على اليمين)
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(5, 5, 5, 5)

        right_label = QLabel("النص الأصلي")
        right_label.setStyleSheet(
            f"color: {COLORS['fg_primary']}; font-weight: bold; font-size: 12px; padding: 5px;"
        )
        right_layout.addWidget(right_label)

        self.text_input = QTextEdit()
        self.text_input.setStyleSheet(f"""
            QTextEdit {{
                background-color: {COLORS["bg_tertiary"]};
                color: {COLORS["fg_primary"]};
                border: 1px solid {COLORS["border"]};
                border-radius: 3px;
                padding: 10px;
                font-size: 14px;
            }}
        """)
        self.text_input.setFont(QFont("Sakkala Majalla", 14))
        right_layout.addWidget(self.text_input)

        splitter.addWidget(right_widget)

        # تقسيم متساوي
        splitter.setSizes([600, 600])

        main_layout.addWidget(splitter)

    def load_corrections_file(self):
        """تحميل ملف التصحيحات من ملف"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختر ملف التصحيحات", "", "ملفات JSON (*.json);;جميع الملفات (*.*)"
        )

        if file_path:
            if self.load_corrections(file_path):
                self.corrections_file_path = file_path
                QMessageBox.information(
                    self, "نجح", f"تم تحميل {len(self.corrections)} تصحيح"
                )

    def import_txt(self):
        """استيراد ملف نصي"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختر ملف نصي", "", "ملفات نصية (*.txt);;جميع الملفات (*.*)"
        )

        if file_path:
            try:
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(file_path, "r", encoding="windows-1256") as f:
                        content = f.read()

                self.text_input.setPlainText(content)
                self.statusBar().showMessage("تم استيراد الملف بنجاح", 3000)
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل في قراءة الملف:\n{str(e)}")

    def import_docx(self):
        """استيراد ملف Word"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(
                self,
                "غير متاح",
                "مكتبة python-docx غير مثبتة.\nقم بتثبيتها باستخدام:\npip install python-docx",
            )
            return

        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختر ملف Word", "", "ملفات Word (*.docx);;جميع الملفات (*.*)"
        )

        if file_path:
            try:
                doc = Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

                self.text_input.setPlainText(content)
                self.statusBar().showMessage("تم استيراد الملف بنجاح", 3000)
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل في قراءة الملف:\n{str(e)}")

    def find_errors(self, text):
        """البحث عن الأخطاء في النص"""
        errors = []

        if not self.wrong_words:
            return errors

        for wrong_word in self.wrong_words:
            pattern = r"\b" + re.escape(wrong_word) + r"\b"
            matches = re.finditer(pattern, text, re.IGNORECASE)

            for match in matches:
                start_pos = match.start()
                end_pos = match.end()

                # الحصول على السياق
                context_start = max(0, start_pos - 30)
                context_end = min(len(text), end_pos + 30)
                context = text[context_start:context_end]

                errors.append(
                    {
                        "word": match.group(),
                        "correct": self.corrections[wrong_word],
                        "position": start_pos,
                        "end_position": end_pos,
                        "context": context,
                    }
                )

        return errors

    def create_corrected_text_with_inline_diffs(self):
        """إنشاء النص المصحح مع inline diffs بأسلوب GitHub"""
        if not self.errors:
            self.corrected_text.clear()
            return

        # الحصول على النص الأصلي
        text = self.text_input.toPlainText()

        # فرز الأخطاء حسب الموقع (من الآخر للأول لتجنب مشاكل الفهرسة)
        sorted_errors = sorted(self.errors, key=lambda x: x["position"], reverse=True)

        # إنشاء HTML للنص المصحح
        html_parts = [
            """<div style="font-family: 'Sakkala Majalla', 'Arial'; direction: rtl; text-align: right; 
                line-height: 1.8; word-wrap: break-word; white-space: pre-wrap;">"""
        ]

        last_pos = len(text)

        for error in sorted_errors:
            # إضافة النص بعد هذا الخطأ
            html_parts.append(self.escape_html(text[error["end_position"] : last_pos]))

            # إضافة التغيير بأسلوب inline diff - مع مسافة بين الخطأ والتصحيح
            html_parts.append(
                f'<span style="background-color: {COLORS["error_bg"]}; color: {COLORS["error"]}; '
                f"text-decoration: line-through; padding: 2px 4px; border-radius: 2px; "
                f'white-space: nowrap; display: inline-block; margin-left: 4px;">{self.escape_html(error["word"])}</span>'
                f" "
                f'<span style="background-color: {COLORS["success_bg"]}; color: {COLORS["success"]}; '
                f"font-weight: bold; padding: 2px 4px; border-radius: 2px; "
                f'white-space: nowrap; display: inline-block;">{self.escape_html(error["correct"])}</span>'
            )

            last_pos = error["position"]

        # إضافة باقي النص من البداية
        html_parts.append(self.escape_html(text[0:last_pos]))

        html_parts.append("</div>")

        # عكس القائمة لأننا بنينا من الآخر
        html_content = "".join(reversed(html_parts))

        self.corrected_text.setHtml(html_content)

    def escape_html(self, text):
        """تحويل النص لـ HTML آمن مع الحفاظ على التفاف الكلمات الطبيعي"""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    def highlight_error_in_text(self, item):
        """تمييز خطأ معين عند النقر عليه في القائمة"""
        # استخراج الكلمة الخاطئة من النص
        item_text = item.text()
        # النص بصيغة: "❌ كلمة_خاطئة → كلمة_صحيحة (التكرار: عدد)"
        parts = item_text.split("→")
        if len(parts) < 2:
            return

        wrong_word = parts[0].replace("❌", "").strip()
        correct_word = parts[1].split("(")[0].strip()

        # البحث عن الخطأ المطابق
        matching_error = next((e for e in self.errors if e["word"] == wrong_word), None)

        if matching_error:
            # التمرير إلى موقع الخطأ في النص الأصلي
            cursor = self.text_input.textCursor()
            cursor.setPosition(matching_error["position"])
            self.text_input.setTextCursor(cursor)
            self.text_input.ensureCursorVisible()

    def check_text(self):
        """فحص النص وإيجاد الأخطاء"""
        if not self.corrections:
            QMessageBox.warning(
                self,
                "تحذير",
                "لم يتم تحميل ملف التصحيحات.\nيرجى تحميل ملف التصحيحات أولاً.",
            )
            return

        text = self.text_input.toPlainText().strip()

        if not text:
            QMessageBox.warning(self, "تحذير", "لا يوجد نص للفحص")
            return

        # البحث عن الأخطاء
        self.errors = self.find_errors(text)

        # عرض الأخطاء
        self.errors_list.clear()

        if self.errors:
            error_counts = {}
            for error in self.errors:
                word = error["word"]
                if word in error_counts:
                    error_counts[word] += 1
                else:
                    error_counts[word] = 1

            for word, count in error_counts.items():
                error_info = next(e for e in self.errors if e["word"] == word)
                item_text = f"❌ {word} → {error_info['correct']} (التكرار: {count})"
                item = QListWidgetItem(item_text)
                self.errors_list.addItem(item)

            self.stats_label.setText(
                f"عدد الأخطاء: {len(self.errors)} | أنواع مختلفة: {len(error_counts)}"
            )
            self.statusBar().showMessage(f"تم اكتشاف {len(self.errors)} خطأ", 3000)

            # إنشاء النص المصحح مع inline diffs
            self.create_corrected_text_with_inline_diffs()
        else:
            self.stats_label.setText("عدد الأخطاء: 0")
            self.statusBar().showMessage("✅ لم يتم العثور على أخطاء", 3000)
            self.corrected_text.clear()
            QMessageBox.information(
                self, "ممتاز", "لم يتم العثور على أخطاء لغوية في النص!"
            )

    def export_result(self):
        """تصدير النتيجة إلى ملف نصي"""
        if not self.errors:
            QMessageBox.warning(self, "تحذير", "لا توجد أخطاء للتصدير")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "حفظ النتيجة", "", "ملفات نصية (*.txt);;جميع الملفات (*.*)"
        )

        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write("تقرير الأخطاء اللغوية\n")
                    f.write("=" * 50 + "\n\n")

                    error_counts = {}
                    for error in self.errors:
                        word = error["word"]
                        if word in error_counts:
                            error_counts[word] += 1
                        else:
                            error_counts[word] = 1

                    for word, count in error_counts.items():
                        error_info = next(e for e in self.errors if e["word"] == word)
                        f.write(f"الكلمة الخاطئة: {word}\n")
                        f.write(f"الصحيح: {error_info['correct']}\n")
                        f.write(f"التكرار: {count}\n")
                        f.write(f"السياق: {error_info['context']}\n")
                        f.write("-" * 50 + "\n")

                self.statusBar().showMessage("تم حفظ الملف بنجاح", 3000)
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل في حفظ الملف:\n{str(e)}")

    def export_docx_report(self):
        """تصدير تقرير DOCX"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(
                self,
                "غير متاح",
                "مكتبة python-docx غير مثبتة.\nقم بتثبيتها باستخدام:\npip install python-docx",
            )
            return

        if not self.errors:
            QMessageBox.warning(self, "تحذير", "لا توجد أخطاء للتصدير")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "حفظ التقرير", "", "ملفات Word (*.docx);;جميع الملفات (*.*)"
        )

        if file_path:
            try:
                doc = Document()

                # العنوان
                title = doc.add_heading("تقرير تصحيح الأخطاء اللغوية", 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # معلومات
                doc.add_paragraph(
                    f"تاريخ المعالجة: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                )
                doc.add_paragraph(f"عدد الأخطاء: {len(self.errors)}")

                # جدول الأخطاء
                if self.errors:
                    doc.add_heading("الأخطاء المكتشفة", level=1)
                    errors_table = doc.add_table(rows=1, cols=4)
                    errors_table.style = "Table Grid"

                    headers = ["التكرار", "الأصوب", "الكلمة الخاطئة", "السياق"]
                    for i, header in enumerate(headers):
                        cell = errors_table.rows[0].cells[i]
                        cell.text = header
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cell.paragraphs[0].runs[0].font.bold = True

                    error_counts = {}
                    for error in self.errors:
                        word = error["word"]
                        if word in error_counts:
                            error_counts[word] += 1
                        else:
                            error_counts[word] = 1

                    for word, count in error_counts.items():
                        error_info = next(e for e in self.errors if e["word"] == word)
                        row = errors_table.add_row()

                        row.cells[0].text = str(count)
                        row.cells[1].text = error_info["correct"]
                        row.cells[2].text = word
                        row.cells[3].text = error_info["context"]

                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                doc.save(file_path)
                self.statusBar().showMessage("تم حفظ التقرير بنجاح", 3000)
            except Exception as e:
                QMessageBox.critical(self, "خطأ", f"فشل في حفظ التقرير:\n{str(e)}")

    def clear_all(self):
        """مسح جميع النصوص"""
        reply = QMessageBox.question(
            self,
            "تأكيد",
            "هل تريد مسح جميع النصوص؟",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.text_input.clear()
            self.errors_list.clear()
            self.corrected_text.clear()
            self.errors = []
            self.stats_label.setText("عدد الأخطاء: 0")
            self.statusBar().showMessage("تم المسح", 2000)


def check_dependencies():
    """فحص المكتبات المطلوبة"""
    missing = []

    if not PYQT6_AVAILABLE:
        missing.append("PyQt6")

    if not DOCX_AVAILABLE:
        missing.append("python-docx")

    if missing:
        print("❌ المكتبات التالية غير مثبتة:")
        for lib in missing:
            print(f"   - {lib}")
        print("\nقم بتثبيتها باستخدام:")
        for lib in missing:
            if lib == "PyQt6":
                print(f"   pip install {lib}")
            else:
                print(f"   pip install {lib}")
        return False

    return True


def main():
    if not check_dependencies():
        sys.exit(1)

    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    window = StyleCheckerApp()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
